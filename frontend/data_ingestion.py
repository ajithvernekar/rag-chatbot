import os
import logging, time
from dotenv import load_dotenv
from langchain_pinecone import PineconeVectorStore
from langchain_openai import OpenAIEmbeddings
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pinecone import Pinecone, ServerlessSpec
from langchain.schema import Document
from docx import Document as DocxDocument

# Load environment variables
load_dotenv()
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_ENVIRONMENT = os.getenv("PINECONE_ENVIRONMENT")
INDEX_NAME = os.getenv("INDEX_NAME")

# Configure logging
logging.basicConfig(level=logging.INFO)

def initialize_pinecone():
    """Initialize Pinecone client."""
    pinecone = Pinecone(PINECONE_API_KEY)
    spec = ServerlessSpec(cloud='aws', region=PINECONE_ENVIRONMENT)
    return pinecone, spec

def delete_and_create_index(pinecone, spec, index_name):
    """Delete the existing Pinecone index and create a new one."""
    if pinecone.has_index(index_name):
        logging.info(f"Deleting existing Pinecone index: {index_name}")
        pinecone.delete_index(index_name)
    logging.info(f"Creating new Pinecone index: {index_name}")
    pinecone.create_index(
        name=index_name,
        dimension=1536,  # OpenAI embedding dimensions
        metric="cosine",
        spec=spec
    )

def split_text_into_chunks(text: str):
    """Split long text into smaller chunks."""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return text_splitter.split_text(text)

def process_and_index_file(file_path: str, openai_api_key: str):
    """Load, process, and index documents into Pinecone."""
    if file_path.endswith(".pdf"):
        loader = PyPDFLoader(file_path)
        documents = loader.load()
    elif file_path.endswith(".docx"):
        logging.info(f"Loading .docx file: {file_path}")
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        documents = [Document(page_content=text, metadata={"source": file_path})]
    elif file_path.endswith(".txt"):
        loader = TextLoader(file_path)
        documents = loader.load()
    else:
        raise ValueError("Unsupported file format")
    
    logging.info(f"Loaded {len(documents)} documents from {file_path}")

    processed_documents = []

    for doc in documents:
        chunks = split_text_into_chunks(doc.page_content)
        for chunk in chunks:
            processed_documents.append(Document(page_content=chunk, metadata=doc.metadata))

    # Initialize Pinecone
    pinecone, spec = initialize_pinecone()

    # Delete and recreate the index
    delete_and_create_index(pinecone, spec, INDEX_NAME)

    # Add processed documents to the vector store
    logging.info("Generating embeddings for documents...")

    try:
        vectorstore = PineconeVectorStore(
            index_name=INDEX_NAME,
            embedding=OpenAIEmbeddings(model="text-embedding-ada-002", openai_api_key=openai_api_key)
        )
        logging.info("PineconeVectorStore initialized successfully.")
    except Exception as e:
        logging.error(f"Error initializing PineconeVectorStore: {str(e)}")
        raise

    logging.info("Adding documents to Pinecone...")
    
    try:
        vectorstore.add_documents(processed_documents)
        logging.info("Documents added successfully.")
    except Exception as e:
        logging.error(f"Error adding documents to Pinecone: {str(e)}")
        raise

if __name__ == "__main__":
    # Example usage
    file_path = "data/Atomic_habits.pdf"  # Update this path accordingly  
    process_and_index_file(file_path, os.getenv("OPENAI_API_KEY"))